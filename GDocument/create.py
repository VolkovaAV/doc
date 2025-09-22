from docx import Document
import config
import os
import pandas as pd
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

script_path = os.path.dirname(os.path.abspath(__file__))

def remove_table_borders(table):
    """
    Убирает все границы у таблицы, корректно работая с oxml-узлами.
    """
    tbl = table._tbl

    # получить/создать <w:tblPr>
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # получить/создать <w:tblBorders>
    tblBorders = None
    for child in tblPr:
        if child.tag == qn('w:tblBorders'):
            tblBorders = child
            break
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # очистить ранее заданные границы
    for child in list(tblBorders):
        tblBorders.remove(child)

    # задать все стороны как 'nil' (нет границ)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'nil')
        tblBorders.append(elem)


def make_row_bold(row):
    """Делает весь текст в строке жирным."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

def set_cell_border(cell, **kwargs):
    """
    Устанавливает границы ячейки.
    Пример вызова: set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})
    Доступные стороны: top, bottom, start(=left), end(=right), insideH, insideV
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Найдём существующий <w:tcBorders> или создадим новый
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            element = tcBorders.find(qn(f'w:{edge}'))
            if element is None:
                element = OxmlElement(f'w:{edge}')
                tcBorders.append(element)
            for key in ["val", "sz", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))

def create_act_template_doc(filename):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)     # верхнее
    section.bottom_margin = Cm(2)  # нижнее
    section.left_margin = Cm(3)    # левое
    section.right_margin = Cm(1.5) # правое

    # Добавим пример 
    # Создаём абзац
    title = doc.add_paragraph("АКТ")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format = title.runs[0].font
    title_format.name = "Times New Roman"
    title_format.size = Pt(11)
    title_format.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # центрирование всего параграфа
    
    # --- Первая часть текста ---
    run1 = p.add_run("об оказании услуг по публичной оферте на оказание услуг по организации участия в ")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(11)
    run1.font.bold = True

    # --- МЕРОПРИЯТИЕ ---
    run2 = p.add_run("НАЗВАНИЕ МЕРОПРИЯТИЯ")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW  # жёлтое выделение

    doc.add_paragraph()

    # --- ТЕКСТ ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # центрирование всего параграфа

    run03 = p1.add_run('{{ LAST_NAME }} {{ FIRST_NAME }} {{ MIDDLE_NAME }}')
    run03.font.name = "Times New Roman"
    run03.font.size = Pt(11)
    run03.underline = True


    run3 = p1.add_run(', именуе{{ SEX }} в дальнейшем «Заказчик», с одной стороны, и МЦФПИН в лице ректора Евтушенко Андрея Александровича, действующего на основании Устава, именуемый в дальнейшем «Исполнитель», составили настоящий акт в подтверждение того, что Исполнителем оказаны услуги по организации участия Заказчика в ')
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(11)
    
    run4 = p1.add_run("НАЗВАНИЕ МЕРОПРИЯТИЯ")
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(11)
    run4.font.highlight_color = WD_COLOR_INDEX.YELLOW  # жёлтое выделение

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # центрирование всего параграфа

    run5 = p2.add_run('Стоимость оказанных услуг {{ SUMM }} руб. 00 коп. ({{ SUMM_NAME }} рублей 00 копеек). НДС не облагается на основании п. 1 ст. 145 НК РФ.')
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(11)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # центрирование всего параграфа

    run6 = p3.add_run('Вышеперечисленные услуги выполнены полностью и в срок. Заказчик претензий по объему, качеству и срокам оказания услуг не имеет.')
    run6.font.name = "Times New Roman"
    run6.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Таблица 2×4, автоподбор ширины
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    remove_table_borders(table)

    # Заполнение + первая строка жирная
    data = [
        ("ЗАКАЗЧИК", "ИСПОЛНИТЕЛЬ"),
        (" ", "МЦФПИН"),
        (" ", "Ректор"),
        ("___________/{{ F_NAME }}{{ M_NAME }}{{ LAST_NAME }}", "___________/А.А.Евтушенко"),
    ]
    for row_idx, (c1, c2) in enumerate(data):
        row = table.rows[row_idx]
        row.cells[0].text, row.cells[1].text = c1, c2
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
        if row_idx == 0:  # заголовок жирный
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
    doc.save(filename)
    return f"Создан документ-шаблон '{filename}'."

def create_bill_template_doc(filename="complex_table.docx"):
    doc = Document()

    title = doc.add_paragraph("СЧЁТ")
    # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format = title.runs[0].font
    title_format.name = "Times New Roman"
    title_format.size = Pt(11)
    title_format.bold = True

    doc.add_paragraph()

    # Таблица 3 столбца × 4 строки
    table = doc.add_table(rows=17, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Настраиваем ширину каждого столбца
    widths = [Cm(4.5), Cm(0.5), Cm(4.5), Cm(0.5),  Cm(6) ]  # ширины столбцов
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    

    # Удобный доступ к ячейкам (0-индексация)
    c = lambda r, col: table.cell(r, col)

    # --- Объединения ---
    c(2, 0).merge(c(2, 2))
    c(3, 0).merge(c(3, 2))
    c(4, 0).merge(c(4, 2))
    c(5, 0).merge(c(5, 2))
    c(8, 0).merge(c(8, 4))
    c(9, 0).merge(c(9, 4))
    c(10, 0).merge(c(10, 4))
    c(11, 0).merge(c(11, 4))
    c(12, 0).merge(c(12, 4))
    c(13, 0).merge(c(13, 4))
    c(14, 0).merge(c(14, 4))
    c(15, 0).merge(c(15, 4))
    c(16, 0).merge(c(16, 4))
    c(0, 4).merge(c(7, 4))


    # --- Заполнение текстом ---
    # Строка 1 (без объединений в 1–2 столбцах)
    c(0, 0).text = "МЦФПИН"
    c(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    

    c(1, 0).text = " "
    p = c(1, 0).paragraphs[0]
    run = p.add_run("Наименование получателя платежа")
    run.font.italic = True
    run.font.size = Pt(10)

    c(0, 2).text = "5260054053/526001001"
    c(1, 2).text = " "
    p = c(1, 2).paragraphs[0]
    run = p.add_run("ИНН/КПП")
    run.font.italic = True
    run.font.size = Pt(10)
    

    # Строка 2 (ячейка 1–2 объединена, писать в левую верхнюю)
    c(2, 0).text = "40703810942000000672"
    c(3, 0).text = " "
    p = c(3, 0).paragraphs[0]
    run = p.add_run("Номер счета получателя платежа")
    run.font.italic = True
    run.font.size = Pt(10)
    
    c(4, 0).text = "ВОЛГО-ВЯТСКИЙ БАНК ПАО СБЕРБАНК"

    c(5, 0).text = " "
    p = c(5, 0).paragraphs[0]
    run = p.add_run("Банк получателя платежа")
    run.font.italic = True
    run.font.size = Pt(10)

    c(6, 0).text = "042202603"
    c(6, 2).text = "30101810900000000603"

    c(7, 0).text = " "
    p = c(7, 0).paragraphs[0]
    run = p.add_run("БИК банка")
    run.font.italic = True
    run.font.size = Pt(10)

    c(7, 2).text = " "
    p = c(7, 2).paragraphs[0]
    run = p.add_run("Корр. счет банка")
    run.font.italic = True
    run.font.size = Pt(10)

    c(8, 0).text = "Оплата регистрационного взноса за участие в "

    c(9, 0).text = " "
    p = c(9, 0).paragraphs[0]
    run = p.add_run("НАЗВАНИЕ МЕРОПРИЯТИЯ")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # жёлтое выделение

    c(10, 0).text = " "
    p = c(10, 0).paragraphs[0]
    run = p.add_run("Назначение платежа")
    run.font.italic = True
    run.font.size = Pt(10)

    c(11, 0).text = " "
    p = c(11, 0).paragraphs[0]
    run = p.add_run("ДАТЫ ПРОВЕДЕНИЯ")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # жёлтое выделение

    c(12, 0).text = " "
    p = c(12, 0).paragraphs[0]
    run = p.add_run("Время проведения мероприятия")
    run.font.italic = True
    run.font.size = Pt(10)

    c(13, 0).text = "{{ LAST_NAME }} {{ FIRST_NAME }} {{ MIDDLE_NAME }}"

    c(14, 0).text = " "
    p = c(14, 0).paragraphs[0]
    run = p.add_run("Плательщик")
    run.font.italic = True
    run.font.size = Pt(10)

    c(15, 0).text = "{{ SUMM }} руб. 00 коп."

    c(16, 0).text = " "
    p = c(16, 0).paragraphs[0]
    run = p.add_run("Сумма платежа")
    run.font.italic = True
    run.font.size = Pt(10)

    c(0, 4).text = "{{ FILENAME }}"
    


    # --- Форматирование: Times New Roman, 12 pt; выравнивания ---
    for row in table.rows:
        for cell in row.cells:
            # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(4)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    # run.font.size = Pt(11)
                
    # --- Устанавливаем только нижнюю границу для ячейки (0,0) ---

    target_cells = [c(0, 0), c(0, 2), c(2, 0), c(4, 0), c(6, 0), c(6, 2), c(8, 0), c(9, 0), c(11, 0), c(13, 0), c(15, 0)]
    for _c in target_cells:
        set_cell_border(
            _c,
            bottom={"val": "single", "sz": 5, "color": "000000"}  # чёрная линия толщиной 12
        )


    doc.add_paragraph()

    table1= doc.add_table(rows=1, cols=3)
    c = lambda r, col: table1.cell(r, col)

    c(0, 0).text = 'Ректор'
    c(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    
    c(0, 1).text = ' '
    par = c(0, 1).paragraphs[0]
    run = par.add_run()
    run.add_picture(f'{script_path}/ver_02.png')

    c(0, 2).text = 'А.А.Евтушенко'
    c(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    
    doc.save(filename)
    return f"Создан документ-шаблон '{filename}'."

def create_excel_with_columns(filename, columns):
    """
    Создаёт пустой Excel-файл с нужными столбцамиП, если такого файла ещё нет.
    """
    df = pd.DataFrame(columns=columns)
    df.to_excel(filename, sheet_name="Лист1", index=False)
    return f"Файл '{filename}' успешно создан."


def create_all_templates():
    if not os.path.exists(config.TEMP_FOLDER_NAME):
        os.makedirs(config.TEMP_FOLDER_NAME)
        print(os.path.exists(config.TEMP_FOLDER_NAME))

    if not os.path.isfile(f'{config.TEMP_FOLDER_NAME}/act.docx'):
        res1 =create_act_template_doc(f'{config.TEMP_FOLDER_NAME}/act.docx') + '\n'
    else:
        res1 = f'Файл {config.TEMP_FOLDER_NAME}/act.docx существует' + '\n'

    if not os.path.isfile(f'{config.TEMP_FOLDER_NAME}/bill.docx'):
        res2 =create_bill_template_doc(f'{config.TEMP_FOLDER_NAME}/bill.docx') + '\n'
    else:
        res2 = f'Файл {config.TEMP_FOLDER_NAME}/bill.docx существует' + '\n'

    if not os.path.isfile(config.TB_NAME):
        res3 = create_excel_with_columns(config.TB_NAME, config.STD_COL_NAME) + '\n'
    else: res3 = f'Файл {config.TB_NAME} существует' + '\n'

    return res1+res2+res3
